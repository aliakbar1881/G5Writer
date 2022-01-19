from filehandeling import main as filemain
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.shared import Inches


def main():
    word = list()
    type_ = list()
    example = list()
    explain = list()
    for key, value in filemain().items():
        def Word_Handeling(key, value):
            if len(word) == 2:
                writing(word, type_, explain, example)

            else:
                word.append(key)
                type_.append(value.type_)
                example.append(value.example)
                explain.append(value.explain)


        Word_Handeling(key, value)


def writing(word, type_, explain, example):
    path = '/home/aliakbar/personal/python/practice/writing-G5/word/mypart.docx'

    document = Document(path)
    section = document.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(5.83)
    section.page_height = Inches(4.13)
    for i in range(2):
        if i == 0:
            for i in range(2):
                paragraph = document.add_paragraph(word[i])
                paragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                font = run.font
                font.name = 'Calibri'
                font.size = Pt(12)
            document.add_page_break()
        else:
            for i in range(2):
                paragraph = document.add_paragraph(explain[i])
                paragraph = document.add_paragraph(example[i])
                pragraph_format = paragraph.paragraph_format
                paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = paragraph.add_run()
                font = run.font
                font.name = 'Calibri'
                font.size = Pt(9)
            document.add_page_break()



    document.save(path)


if __name__ == '__main__':
    main()
